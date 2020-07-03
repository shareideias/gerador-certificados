import os
import sys

import postgresql

from shutil import copyfile
from shutil import copytree
from shutil import make_archive
from shutil import rmtree
from shutil import move

from docx import Document
from docx2pdf import convert

from Participante import Participante

def main():
    try:
        id_curso = sys.argv[1]
    except:
        print("Curso não especificado")
        return

    #Conexão com o BD 
    try:
        db = postgresql.open('pq://postgres:4223@localhost:5432/shareideias')
        sql = "SELECT par.nome, par.curso1_id AS curso_id, cur.nome, par.data_inscricao_c1 AS data_inscricao, CURRENT_TIMESTAMP AS data_cert FROM sisins_participante AS par JOIN sisins_curso AS cur ON par.curso1_id = cur.id  WHERE (par.curso1_id = {0} AND par.resultado_c1 = 1) UNION SELECT par.nome, par.curso2_id AS curso_id, cur.nome, par.data_inscricao_c2 AS data_inscricao, CURRENT_TIMESTAMP AS data_cert FROM sisins_participante AS par JOIN sisins_curso AS cur ON par.curso2_id = cur.id  WHERE (par.curso2_id = {0} AND par.resultado_c2 = 1)".format(id_curso)
        query = db.prepare(sql)
    except:
        print("Erro na conexão com o BD:", sys.exc_info())
        return

    certPath = os.path.dirname(__file__) + "/certificados"
    zipPath = os.path.dirname(__file__) + '/Certificados'
    pathBase = os.path.dirname(__file__) + "/base/aluno.docx"
    pathCopia = certPath + "/aluno.docx"
    public = os.path.dirname(__file__) + "/../../public"

    #Tenta realizar
    try:
        #Cria o diretório dos certificados
        rmtree(certPath, ignore_errors=True)
        os.mkdir(certPath)

        #Repete para todos os alunos
        for row in query:
            aluno = Participante(row)
            #Gera o certificado de cada aluno
            geraCertificado(aluno, pathBase)

        #Zipa todos
        make_archive(zipPath, 'zip', certPath)
        move(os.path.join(os.path.dirname(__file__), 'Certificados.zip'), os.path.join(public, 'Certificados.zip')) 
    except:
        print("Erro na geração de arquivos:", repr(sys.exc_info()[0]), repr(sys.exc_info()[1]), sys.exc_info()[2].print_exc())
        return
    finally:
        #Exclui os certificados
        rmtree(certPath, ignore_errors=True)
        pass

# Deve pegar editar o modelo em memória e colocar no padrão, zipar e excluir o xml
def geraCertificado(aluno, pathBase):
    #DOCX
    pathNovo = os.path.dirname(__file__) + "/certificados/{0}.docx".format(aluno.nome)
    document = Document(pathBase)

    textos_ini = document.paragraphs[0].runs
    textos_data = document.paragraphs[5].runs
    idx = [8, 12, 14]
    for i in idx:
        textos_ini[i].text = textos_ini[i].text.replace("(NOME)", aluno.nome.upper()).replace("(CURSO)", aluno.curso.upper()).replace("(SEMESTRE)", aluno.getSemestre().upper()).replace("(ANO)", str(aluno.data_inscricao.year)).replace("(ANOC)", str(aluno.data_certificado.year)).replace("(MES)", aluno.getMes().upper()).replace("(DIA)", str(aluno.data_certificado.day))
    textos_data[0].text = textos_data[0].text.replace("(NOME)", aluno.nome.upper()).replace("(CURSO)", aluno.curso.upper()).replace("(SEMESTRE)", aluno.getSemestre().upper()).replace("(ANO)", str(aluno.data_inscricao.year)).replace("(ANOC)", str(aluno.data_certificado.year)).replace("(MES)", aluno.getMes().upper()).replace("(DIA)", str(aluno.data_certificado.day))

    document.save(pathNovo)

    #PDF
    convert(pathNovo)
    os.remove(pathNovo)

if __name__ == "__main__":
    main()